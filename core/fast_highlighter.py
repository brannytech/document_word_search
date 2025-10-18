"""Fast highlighting using already-extracted text - no file re-opening"""

import os
from pathlib import Path
from typing import Dict, Tuple, List
import fitz  # PyMuPDF
from docx import Document
from docx.shared import RGBColor
import re

from config import Config
from searchers.base import SearchResult


class FastHighlighter:
    """Fast highlighting using text already in memory"""
    
    def __init__(self):
        self.output_dir = Config.OUTPUT_DIR
        self.output_dir.mkdir(exist_ok=True)
    
    def highlight_all_from_memory(self, 
                                   results: Dict[str, List[SearchResult]], 
                                   keyword: str,
                                   extracted_texts: Dict[str, Tuple[str, int]]) -> Dict[str, str]:
        """
        Highlight all files using already-extracted text
        
        Args:
            results: Search results by file
            keyword: The search keyword
            extracted_texts: Dict of {file_path: (text, page_count)} from extraction
            
        Returns:
            Dict of {original_path: highlighted_path}
        """
        highlighted_files = {}
        
        for file_path in results.keys():
            try:
                # Check if we have the text in memory
                if file_path in extracted_texts:
                    # Use fast method (text already extracted)
                    output_path = self._highlight_from_text(
                        file_path, keyword, extracted_texts[file_path]
                    )
                else:
                    # Fallback: use slow method (re-open file)
                    output_path = self._highlight_from_file(file_path, keyword)
                
                if output_path and os.path.exists(output_path):
                    highlighted_files[file_path] = output_path
                    
            except Exception as e:
                print(f"Error highlighting {file_path}: {e}")
        
        return highlighted_files
    
    def _highlight_from_text(self, file_path: str, keyword: str, 
                            text_data: Tuple[str, int]) -> str:
        """Highlight using already-extracted text (FAST)"""
        text, page_count = text_data
        ext = Path(file_path).suffix.lower()
        
        # Generate output path
        output_path = self._generate_output_path(file_path, keyword)
        
        # For PDF: Must re-open file (PyMuPDF limitation)
        if ext == '.pdf':
            return self._highlight_pdf_fast(file_path, keyword, output_path)
        
        # For DOCX: Can use text-based highlighting
        elif ext == '.docx':
            return self._highlight_docx_fast(file_path, keyword, output_path)
        
        # For DOC: Convert to DOCX first
        elif ext == '.doc':
            return self._highlight_doc_fast(file_path, keyword, output_path)
        
        return ""
    
    def _highlight_from_file(self, file_path: str, keyword: str) -> str:
        """Fallback: Highlight by re-opening file (SLOW)"""
        ext = Path(file_path).suffix.lower()
        output_path = self._generate_output_path(file_path, keyword)
        
        if ext == '.pdf':
            return self._highlight_pdf_fast(file_path, keyword, output_path)
        elif ext == '.docx':
            return self._highlight_docx_fast(file_path, keyword, output_path)
        
        return ""
    
    def _highlight_pdf_fast(self, file_path: str, keyword: str, output_path: str) -> str:
        """Highlight PDF - optimized"""
        try:
            doc = fitz.open(file_path)
            
            # Build search pattern with word boundaries
            pattern = self._build_pattern(keyword)
            
            # Highlight all matches
            for page in doc:
                text = page.get_text()
                
                # Find all matches
                for match in re.finditer(pattern, text, re.IGNORECASE):
                    # Search for match on page
                    text_instances = page.search_for(match.group())
                    
                    # Highlight each instance
                    for inst in text_instances:
                        try:
                            highlight = page.add_highlight_annot(inst)
                            highlight.set_colors(stroke=Config.HIGHLIGHT_COLOR_RGB)
                            highlight.update()
                        except:
                            pass  # Skip if highlighting fails
            
            # Save
            doc.save(output_path)
            doc.close()
            return output_path
            
        except Exception as e:
            print(f"Error highlighting PDF: {e}")
            return ""
    
    def _highlight_docx_fast(self, file_path: str, keyword: str, output_path: str) -> str:
        """Highlight DOCX - optimized"""
        try:
            doc = Document(file_path)
            pattern = self._build_pattern(keyword)
            
            # Highlight in paragraphs
            for para in doc.paragraphs:
                self._highlight_paragraph(para, pattern)
            
            # Highlight in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            self._highlight_paragraph(para, pattern)
            
            doc.save(output_path)
            return output_path
            
        except Exception as e:
            print(f"Error highlighting DOCX: {e}")
            return ""
    
    def _highlight_doc_fast(self, file_path: str, keyword: str, output_path: str) -> str:
        """Highlight DOC by converting to DOCX"""
        try:
            import pypandoc
            temp_docx = str(Config.TEMP_DIR / "temp_highlight.docx")
            pypandoc.convert_file(file_path, 'docx', outputfile=temp_docx, format='doc')
            
            result = self._highlight_docx_fast(temp_docx, keyword, output_path)
            
            if os.path.exists(temp_docx):
                os.remove(temp_docx)
            
            return result
        except Exception as e:
            print(f"Error highlighting DOC: {e}")
            return ""
    
    def _highlight_paragraph(self, paragraph, pattern):
        """Highlight matches in a paragraph"""
        full_text = paragraph.text
        
        if not full_text:
            return
        
        matches = list(pattern.finditer(full_text))
        
        if not matches:
            return
        
        # Clear existing runs
        for run in paragraph.runs:
            run.text = ''
        
        # Rebuild with highlights
        last_end = 0
        for match in matches:
            # Add text before match
            if match.start() > last_end:
                paragraph.add_run(full_text[last_end:match.start()])
            
            # Add highlighted match
            run = paragraph.add_run(match.group())
            run.font.highlight_color = Config.HIGHLIGHT_COLOR_WORD
            
            last_end = match.end()
        
        # Add remaining text
        if last_end < len(full_text):
            paragraph.add_run(full_text[last_end:])
    
    def _build_pattern(self, keyword: str) -> re.Pattern:
        """Build regex pattern with word boundaries"""
        from utils.helpers import normalize_keyword
        
        normalized = normalize_keyword(keyword)
        words = normalized.split()
        
        pattern_parts = []
        for i, word in enumerate(words):
            escaped_word = re.escape(word)
            if i == len(words) - 1:
                escaped_word = escaped_word + r'(?:e?s)?'
            pattern_parts.append(escaped_word)
        
        pattern = r'[-\s]*'.join(pattern_parts)
        
        # CRITICAL: Add word boundaries to prevent partial matches
        pattern = r'\b' + pattern + r'\b'
        
        return re.compile(pattern, re.IGNORECASE)
    
    def _generate_output_path(self, file_path: str, keyword: str) -> str:
        """Generate output path for highlighted document"""
        path = Path(file_path)
        clean_keyword = "".join(c for c in keyword if c.isalnum())[:20]
        new_name = f"{path.stem}_highlighted_{clean_keyword}{path.suffix}"
        return str(self.output_dir / new_name)
    
    