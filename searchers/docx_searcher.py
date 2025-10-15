"""DOCX document searcher with highlighting"""

import os
from typing import List
from docx import Document
from .base import BaseSearcher, SearchResult
from utils.helpers import create_sentence_context
from config import Config


class DOCXSearcher(BaseSearcher):
    """Search and highlight DOCX documents"""
    
    def search(self, file_path: str, keyword: str, 
               case_sensitive: bool = False,
               whole_word: bool = False) -> List[SearchResult]:
        """Search for keyword in DOCX with fuzzy matching"""
        results = []
        
        if self.stop_search:
            return results
        
        try:
            doc = Document(file_path)
            
            # Extract all text first for better performance
            full_text = '\n'.join([para.text for para in doc.paragraphs])
            
            if not full_text:
                return results
            
            # Use fuzzy pattern
            pattern = self._build_fuzzy_pattern(keyword, case_sensitive, whole_word)
            
            # Find all matches in full text
            for match in pattern.finditer(full_text):
                if self.stop_search:
                    break
                
                # Estimate page number
                match_position = match.start()
                page_num = (match_position // Config.CHARS_PER_PAGE_ESTIMATE) + 1
                
                context, rel_start, rel_end = create_sentence_context(
                    full_text, match.start(), match.end()
                )
                
                results.append(SearchResult(
                    file_path=file_path,
                    file_name=os.path.basename(file_path),
                    page_number=page_num,
                    context=context,
                    match_start=rel_start,
                    match_end=rel_end,
                    absolute_position=match_position,
                    matched_text=match.group()
                ))
                
        except Exception as e:
            print(f"Error reading DOCX {file_path}: {str(e)}")
            
        return results
    
    def highlight_document(self, file_path: str, keyword: str, 
                          output_path: str, case_sensitive: bool = False) -> bool:
        """Create DOCX with highlighted keywords"""
        try:
            doc = Document(file_path)
            pattern = self._build_fuzzy_pattern(keyword, case_sensitive, False)
            
            for para in doc.paragraphs:
                self._highlight_paragraph(para, pattern)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            self._highlight_paragraph(para, pattern)
            
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"Error highlighting DOCX: {str(e)}")
            return False
    
    def _highlight_paragraph(self, paragraph, pattern):
        """Highlight matches in a paragraph"""
        full_text = paragraph.text
        
        if not full_text:
            return
        
        matches = list(pattern.finditer(full_text))
        
        if not matches:
            return
        
        for run in paragraph.runs:
            run.text = ''
        
        last_end = 0
        
        for match in matches:
            if match.start() > last_end:
                paragraph.add_run(full_text[last_end:match.start()])
            
            run = paragraph.add_run(match.group())
            run.font.highlight_color = Config.HIGHLIGHT_COLOR_WORD
            
            last_end = match.end()
        
        if last_end < len(full_text):
            paragraph.add_run(full_text[last_end:])